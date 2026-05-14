"""
Tests for local parsers (PDF, DOCX, HTML, plain text).
These run in-process (no subprocess pool) so we call parse() directly.
"""
import hashlib
import io

import pytest


class TestPdfParser:

    def test_text_pdf_extracts_content(self):
        """Create minimal in-memory PDF and verify extraction."""
        fitz = pytest.importorskip("fitz")
        from app.parsers.pdf import parse

        # Build a real PDF in memory
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Hello from page one")
        doc2 = fitz.open()
        p2 = doc2.new_page()
        p2.insert_text((72, 72), "Hello from page two")
        doc.insert_pdf(doc2)

        buf = io.BytesIO()
        doc.save(buf)
        data = buf.getvalue()

        result = parse(data)
        assert "Hello from page one" in result.raw_text
        assert result.parse_mode == "pymupdf"
        assert result.char_count > 0
        assert result.checksum == hashlib.sha256(result.raw_text.encode()).hexdigest()

    def test_page_list_populated(self):
        fitz = pytest.importorskip("fitz")
        from app.parsers.pdf import parse

        doc = fitz.open()
        for i in range(3):
            page = doc.new_page()
            page.insert_text((72, 72), f"Page content {i}")
        buf = io.BytesIO()
        doc.save(buf)

        result = parse(buf.getvalue())
        assert len(result.pages) == 3
        assert result.pages[0]["page_number"] == 1

    def test_empty_pdf_returns_empty_text(self):
        fitz = pytest.importorskip("fitz")
        from app.parsers.pdf import parse

        doc = fitz.open()
        doc.new_page()  # blank page
        buf = io.BytesIO()
        doc.save(buf)

        result = parse(buf.getvalue())
        assert result.raw_text == ""
        assert result.pages == []


class TestDocxParser:

    def test_docx_extracts_paragraphs(self):
        docx = pytest.importorskip("docx")
        from docx import Document
        from app.parsers.docx import parse

        doc = Document()
        doc.add_paragraph("First paragraph")
        doc.add_paragraph("Second paragraph")
        buf = io.BytesIO()
        doc.save(buf)

        result = parse(buf.getvalue())
        assert "First paragraph" in result.raw_text
        assert "Second paragraph" in result.raw_text
        assert result.parse_mode == "python_docx"
        assert result.pages == []  # DOCX has no page structure

    def test_empty_docx_returns_empty(self):
        from docx import Document
        from app.parsers.docx import parse

        doc = Document()
        buf = io.BytesIO()
        doc.save(buf)

        result = parse(buf.getvalue())
        assert result.raw_text == ""

    def test_checksum_consistent(self):
        from docx import Document
        from app.parsers.docx import parse

        doc = Document()
        doc.add_paragraph("Test content")
        buf = io.BytesIO()
        doc.save(buf)
        data = buf.getvalue()

        r1 = parse(data)
        r2 = parse(data)
        assert r1.checksum == r2.checksum


class TestHtmlParser:

    def test_strips_script_tags(self):
        from app.parsers.html import parse

        html = b"<html><body><p>Hello</p><script>alert(1)</script></body></html>"
        result = parse(html)
        assert "Hello" in result.raw_text
        assert "alert" not in result.raw_text

    def test_strips_style_tags(self):
        from app.parsers.html import parse

        html = b"<html><body><p>Content</p><style>.x{color:red}</style></body></html>"
        result = parse(html)
        assert "Content" in result.raw_text
        assert "color" not in result.raw_text

    def test_parse_mode(self):
        from app.parsers.html import parse
        result = parse(b"<p>hello</p>")
        assert result.parse_mode == "beautifulsoup4"

    def test_empty_html_returns_empty(self):
        from app.parsers.html import parse
        result = parse(b"<html><body></body></html>")
        assert result.raw_text == ""

    def test_checksum_matches(self):
        from app.parsers.html import parse
        result = parse(b"<p>test content</p>")
        assert result.checksum == hashlib.sha256(result.raw_text.encode()).hexdigest()


class TestPlainParser:

    def test_plain_text_passthrough(self):
        from app.parsers.plain import parse
        data = b"Hello, world. This is plain text."
        result = parse(data)
        assert result.raw_text == "Hello, world. This is plain text."
        assert result.parse_mode == "plain_text"

    def test_utf8_decoding(self):
        from app.parsers.plain import parse
        data = "Arabic: مرحبا".encode("utf-8")
        result = parse(data)
        assert "مرحبا" in result.raw_text

    def test_invalid_bytes_replaced_not_crashed(self):
        from app.parsers.plain import parse
        bad_bytes = b"Hello \xff\xfe World"
        result = parse(bad_bytes)
        assert "Hello" in result.raw_text
        assert "World" in result.raw_text

    def test_char_count_matches(self):
        from app.parsers.plain import parse
        data = b"Some text here"
        result = parse(data)
        assert result.char_count == len(result.raw_text)

    def test_pages_empty_for_plain(self):
        from app.parsers.plain import parse
        result = parse(b"text")
        assert result.pages == []
