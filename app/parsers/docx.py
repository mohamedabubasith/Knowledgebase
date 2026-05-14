import hashlib
import io

from app.models.domain import ParsedDocument


def parse(data: bytes) -> ParsedDocument:
    from docx import Document

    doc = Document(io.BytesIO(data))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    raw_text = "\n".join(paragraphs)

    return ParsedDocument(
        raw_text=raw_text,
        pages=[],
        parse_mode="python_docx",
        char_count=len(raw_text),
        checksum=hashlib.sha256(raw_text.encode()).hexdigest(),
    )
